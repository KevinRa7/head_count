import json
import re
from typing import Dict, List, Any
from docx import Document
import pdfplumber
import os
def read_docx_file(file_path: str) -> str:
    """
    Read content from a .docx file
    """
    try:
        doc = Document(file_path)
        content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text)
        return '\n'.join(content)
    except Exception as e:
        print(f"Error reading .docx file: {e}")
        return ""
import pdfplumber

def read_pdf_file(file_path: str) -> str:
    """
    Extract text from a PDF file
    """
    try:
        text_content = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text_content.append(page.extract_text() or "")
        return "\n".join(text_content)
    except Exception as e:
        print(f"Error reading PDF file: {e}")
        return ""

def parse_wcart_manual_to_json(file_path: str):
    doc = Document(file_path)
    data = {}
    current_section = None
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        
        # detect headings (any Word "Heading" style)
        if paragraph.style.name.startswith("Heading"):
            current_section = text
            data[current_section] = {"content": "", "subsections": {}}
        else:
            if current_section:
                data[current_section]["content"] += text + "\n"
    
    return data


def convert_to_html(text: str) -> str:
    """
    Convert text formatting to HTML with proper list handling
    """
    if not text:
        return ""
    
    # Split into paragraphs
    paragraphs = text.split('\n\n')
    html_paragraphs = []
    
    for paragraph in paragraphs:
        if not paragraph.strip():
            continue
            
        lines = paragraph.split('\n')
        html_lines = []
        current_list = []
        in_list = False
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Check for bullet points
            if line.startswith('- ') or line.startswith('• ') or re.match(r'^\d+\.\s+', line):
                if not in_list:
                    in_list = True
                    current_list = []
                
                # Clean list item
                if line.startswith('- '):
                    item = line[2:].strip()
                elif line.startswith('• '):
                    item = line[2:].strip()
                else:
                    item = re.sub(r'^\d+\.\s+', '', line).strip()
                
                # Convert formatting within list item
                item_html = convert_formatting(item)
                current_list.append(f"<li>{item_html}</li>")
                
            else:
                # End current list
                if in_list and current_list:
                    list_tag = "ol" if re.match(r'^\d+\.', lines[0]) else "ul"
                    html_lines.append(f"<{list_tag}>{''.join(current_list)}</{list_tag}>")
                    current_list = []
                    in_list = False
                
                # Convert regular line
                html_lines.append(convert_formatting(line))
        
        # Add any remaining list
        if in_list and current_list:
            list_tag = "ol" if re.match(r'^\d+\.', lines[0]) else "ul"
            html_lines.append(f"<{list_tag}>{''.join(current_list)}</{list_tag}>")
        
        # Join lines with breaks
        if html_lines:
            html_paragraphs.append('<br>'.join(html_lines))
    
    return '<br><br>'.join(html_paragraphs)

def convert_formatting(text: str) -> str:
    """
    Convert inline formatting to HTML
    """
    # Bold text **text**
    text = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', text)
    
    # Italic text *text*
    text = re.sub(r'\*(.*?)\*', r'<em>\1</em>', text)
    
    # Underlined text
    text = re.sub(r'_(.*?)_', r'<u>\1</u>', text)
    
    # Code/monospace `text`
    text = re.sub(r'`(.*?)`', r'<code>\1</code>', text)
    
    # Arrows and special characters
    text = text.replace('→', '&rarr;')
    text = text.replace('←', '&larr;')
    text = text.replace('↑', '&uarr;')
    text = text.replace('↓', '&darr;')
    
    return text

def main():
    # Read the .docx file
    file_path = 'Copy of WCART USER MANUAL.pdf'
    
    try:
        # First try reading as .docx
        content = read_docx_file(file_path)
        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".docx":
            content = read_docx_file(file_path)
        elif ext == ".pdf":
            content = read_pdf_file(file_path)
        else:
            print("Unsupported file type")
            return
        
        if not content:
            print("No content found in .docx file. Trying alternative approach...")
            # Try reading as plain text with different encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                    print(f"Successfully read with {encoding} encoding")
                    break
                except UnicodeDecodeError:
                    continue
            else:
                print("Could not read the file with any encoding.")
                return
                
    except FileNotFoundError:
        print(f"File '{file_path}' not found.")
        print("Please make sure the file exists in the current directory.")
        return
    
    # Parse the content to JSON
    content = read_pdf_file(file_path)
    wcart_json = parse_wcart_manual_to_json(content)
    
    # Save to JSON file
    output_file = 'wcart_manual1.json'
    with open(output_file, 'w', encoding='utf-8') as json_file:
        json.dump(wcart_json, json_file, indent=2, ensure_ascii=False)
    
    print("✅ WCART manual converted to JSON successfully!")
    print(f"📁 Output file: {output_file}")
    print(f"📊 Sections processed: {len(wcart_json)}")
    
    # Display summary
    total_subsections = 0
    for section, data in wcart_json.items():
        subsections_count = len(data.get("subsections", {}))
        total_subsections += subsections_count
        content_length = len(data.get("content", ""))
        print(f"• {section}: {subsections_count} subsections, {content_length} chars")
    
    print(f"📈 Total subsections: {total_subsections}")

if __name__ == "__main__":
    main()